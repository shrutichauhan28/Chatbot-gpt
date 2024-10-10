from typing import List
from langchain.docstore.document import Document
from langchain.document_loaders import TextLoader
from PyPDF2 import PdfReader
from PIL import Image  # Import PIL for image handling
from pdf2image import convert_from_path
from docx import Document as DocxDocument
import pytesseract
import pandas as pd
import os
import moviepy.editor as mp
# import whisper  # Use Whisper for transcribing audio from videos
from langchain.text_splitter import RecursiveCharacterTextSplitter
from rerank import rank_chunks_with_bm25
import assemblyai as aai

class CleanTextLoader(TextLoader):
    """Load text files and clean text data."""

    def __init__(self, file_path: str, encoding='utf-8', errors='ignore'):
        """Initialize with file path."""
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        """Load text from file path."""
        with open(self.file_path, encoding=self.encoding) as f:
            text = f.read()
        text = self.clean_text(text)
        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

    def clean_text(self, text: str) -> str:
        """Clean up the text content."""
        lines = (line.strip() for line in text.splitlines())
        text = ' '.join(line for line in lines)
        text = ' '.join(text.split())
        text = '\n'.join(line for line in text.splitlines() if line.strip())
        return text
    

# class ImageLoader:
#     """Load image files and extract text using pytesseract."""

#     def __init__(self, file_path: str):
#         self.file_path = file_path

#     def load(self) -> List[Document]:
#         """Extract text from an image file using OCR."""
#         text = ""
#         try:
#             # Open the image file using PIL
#             with Image.open(self.file_path) as img:
#                 # Apply pytesseract OCR to extract text
#                 text = pytesseract.image_to_string(img)
#         except Exception as e:
#             print(f"Error during OCR extraction for image {self.file_path}: {e}")
        
#         # Clean the extracted text
#         text = CleanTextLoader.clean_text(self, text)
#         metadata = {"source": self.file_path}

#         return [Document(page_content=text, metadata=metadata)]


class PDFLoader:
    """Load PDF files, with fallback to OCR for scanned PDFs."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        """Load text from PDF file with fallback to OCR."""
        # Attempt to extract text using PyPDF
        text = self.extract_text_with_pypdf()
        
        # If no text is extracted, fallback to OCR
        if not text.strip():
            text = self.extract_text_with_ocr()
        
        # Clean the extracted text
        text = CleanTextLoader.clean_text(self, text)
        metadata = {"source": self.file_path}
        
        return [Document(page_content=text, metadata=metadata)]

    def extract_text_with_pypdf(self) -> str:
        """Extract text using PyPDF for PDFs with selectable text."""
        text = ""
        try:
            reader = PdfReader(self.file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
        except Exception as e:
            print(f"Error reading PDF with PyPDF: {e}")
        return text

    def extract_text_with_ocr(self) -> str:
        """Extract text using pytesseract for image-based PDFs."""
        text = ""
        try:
            # Use convert_from_path only for PDF files
            pages = convert_from_path(self.file_path)
            for page in pages:
                # Convert each PDF page to an image and apply OCR
                text += pytesseract.image_to_string(page)
        except Exception as e:
            print(f"Error during OCR extraction for PDF {self.file_path}: {e}")
        return text
    
class DocxLoader:
    """Load DOCX files."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        doc = DocxDocument(self.file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        text = CleanTextLoader.clean_text(self, text)
        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

class CSVLoader:
    """Load CSV files."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        df = pd.read_csv(self.file_path)
        text = df.to_string(index=False)
        text = CleanTextLoader.clean_text(self, text)
        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

class ExcelLoader:
    """Load Excel files."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        df = pd.read_excel(self.file_path)
        text = df.to_string(index=False)
        text = CleanTextLoader.clean_text(self, text)
        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

# Set up AssemblyAI API Key
aai.settings.api_key = "3a6b59d26708443a889c231d59a43ef4"

class VideoLoader:
    """Load video files and extract text using AssemblyAI API."""

    def __init__(self, file_path: str, output_folder: str = "../data"):
        self.file_path = file_path
        self.output_folder = output_folder

    def load(self) -> List[Document]:
        """Extract text from video using AssemblyAI and save it in the specified folder."""
        # Convert local video file path to a remote URL or use local file path directly
        text = self.transcribe_video_with_assemblyai(self.file_path)

        # Check if text was successfully extracted
        if not text.strip():
            print(f"No text extracted from video: {self.file_path}")
            return []

        # Clean the text using the clean_text method of CleanTextLoader
        text = CleanTextLoader.clean_text(self, text)
        metadata = {"source": self.file_path}

        # Save text to file
        text_filename = os.path.basename(self.file_path).replace('.mp4', '.txt').replace('.avi', '.txt')
        self.save_text_to_file(text, text_filename)

        # Check if the file was saved correctly
        text_file_path = os.path.join(self.output_folder, text_filename)
        if not os.path.isfile(text_file_path):
            print(f"Failed to save text file: {text_file_path}")
            return []

        return [Document(page_content=text, metadata=metadata)]

    def transcribe_video_with_assemblyai(self, video_file_path: str) -> str:
        """
        Transcribe video to text using AssemblyAI.

        Args:
            video_file_path (str): Path to the local video file.

        Returns:
            str: Transcribed text from the video.
        """
        # Convert video file to audio (since AssemblyAI API needs audio input)
        audio_output_path = os.path.join(self.output_folder, "temp_audio.wav")
        print("Extracting audio from video...")
        try:
            # Step 1: Extract audio from video
            video = mp.VideoFileClip(video_file_path)
            video.audio.write_audiofile(audio_output_path)

            # Step 2: Use AssemblyAI to transcribe the audio file
            print("Uploading audio to AssemblyAI and starting transcription...")
            transcriber = aai.Transcriber()
            transcript = transcriber.transcribe(audio_output_path)

            # Step 3: Remove temporary audio file
            if os.path.exists(audio_output_path):
                os.remove(audio_output_path)
            print("Transcription complete.")
            return transcript.text
        except Exception as e:
            print(f"Error during transcription with AssemblyAI: {e}")
            return ""

    def save_text_to_file(self, text: str, file_name: str):
        """Save the extracted text to a file in the data folder."""
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)

        output_file_path = os.path.join(self.output_folder, file_name)
        try:
            with open(output_file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Text extracted and saved to: {output_file_path}")
        except Exception as e:
            print(f"Error saving text file: {e}")

# Updated function to load different file types including video files
# --- Function to load, split, and chunk the documents ---
def load_n_split(path: str) -> List[Document]:
    """Load files from the specified directory, clean the text, and split them into chunks."""
    documents = []
    for root, dirs, files in os.walk(path):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                # Initialize the appropriate loader based on file type
                if file.endswith('.txt'):
                    loader = CleanTextLoader(file_path)
                elif file.endswith('.pdf'):
                    loader = PDFLoader(file_path)
                elif file.endswith('.docx'):
                    loader = DocxLoader(file_path)
                elif file.endswith('.csv'):
                    loader = CSVLoader(file_path)
                elif file.endswith('.xlsx'):
                    loader = ExcelLoader(file_path)
                # elif file.endswith(('.png', '.jpg', '.jpeg')):
                    # loader = ImageLoader(file_path)  # Use ImageLoader for image files
                elif file.endswith('.mp4') or file.endswith('.avi'):
                    loader = VideoLoader(file_path)
                else:
                    print(f"Unsupported file type: {file}")
                    continue

                # Load the document and add to the list
                docs = loader.load()
                documents.extend(docs)
            except Exception as e:
                print(f"Error loading file {file_path}: {e}")

    # Use RecursiveCharacterTextSplitter to split the documents
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=80)
    split_docs = []
    for doc in documents:
        split_docs.extend(splitter.split_documents([doc]))

    return split_docs


# Function to load, split, and rerank chunks based on a query using BM25
def load_n_split_and_rerank(path: str, query: str) -> List[Document]:
    """
    Load files from the directory, split them into chunks, and rerank them based on the BM25 algorithm.

    Args:
        path (str): The path to the directory containing the files.
        query (str): The query to use for ranking the chunks.

    Returns:
        List[Document]: The reranked chunks based on relevance to the query.
    """
    # Load and split the documents as usual
    documents = load_n_split(path)
    
    # Rerank the chunks using BM25
    reranked_chunks = rank_chunks_with_bm25(documents, query)
    
    # Optionally return only the reranked chunks (without the BM25 score)
    return [chunk for chunk, score in reranked_chunks]